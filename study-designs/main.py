from docx import Document

def docx_to_md(input_file, output_file):
    doc = Document(input_file)
    md_lines = []

    for para in doc.paragraphs:
        style = para.style.name

        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue

        # Headings
        if style.startswith("Heading"):
            level = int(style.replace("Heading ", ""))
            md_lines.append("#" * level + " " + text)

        # Lists
        elif "List Bullet" in style:
            md_lines.append(f"- {text}")
        elif "List Number" in style:
            md_lines.append(f"1. {text}")

        # Normal paragraph
        else:
            md_lines.append(text)

    # Tables
    for table in doc.tables:
        md_lines.append("")
        for i, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(row_text) + " |")

            # Header separator
            if i == 0:
                md_lines.append("| " + " | ".join(["---"] * len(row.cells)) + " |")
        md_lines.append("")

    with open(output_file, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))


if __name__ == "__main__":
    # get all in pwd
    import os
    for filename in os.listdir("."):
        if filename.endswith(".docx"):
            output_filename = filename.replace(".docx", ".md")
            docx_to_md(filename, output_filename)
            print(f"Converted {filename} to {output_filename}")